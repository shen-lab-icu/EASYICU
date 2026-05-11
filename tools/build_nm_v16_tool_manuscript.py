#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import json
import re
import shutil
import subprocess
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document

ARM_ORDER = ["aware", "aware_no_pref", "naive_with_pref", "naive"]
ARM_LABEL = {
    "aware": "ICU-aware + preferences",
    "aware_no_pref": "ICU-aware only",
    "naive_with_pref": "Generic context + preferences",
    "naive": "Generic context only",
}
TASK_LABEL = {
    "t01_table_one_descriptive": "baseline descriptive summary",
    "t02_outcome_incidence_strata": "outcome incidence by severity strata",
    "t03_severity_score_correlation": "severity-score correlation",
    "t04_lactate_mortality_association": "lactate-mortality association",
    "t05_kdigo_renal_sensitivity": "KDIGO renal sensitivity",
    "t06_shock_phenotype_clustering": "shock phenotype clustering",
    "t07_mortality_prediction_auroc": "mortality prediction performance",
    "t08_vaso_selection_bias_audit": "vasopressor selection-bias audit",
    "t09_sofa_zero_artefact_audit": "SOFA-zero artefact audit",
    "t10_complete_case_robustness": "complete-case robustness",
    "t11_los_distribution_descriptive": "ICU length-of-stay distribution",
    "t12_age_stratified_mortality": "age-stratified mortality",
    "t13_admission_vital_summary": "admission vital-summary profile",
    "t14_creatinine_trajectory_kdigo": "creatinine trajectory and KDIGO analysis",
    "t15_norepinephrine_dose_response": "norepinephrine-equivalent dose-response analysis",
}
DIFFICULTY = {
    "t01_table_one_descriptive": "descriptive",
    "t02_outcome_incidence_strata": "descriptive",
    "t11_los_distribution_descriptive": "descriptive",
    "t12_age_stratified_mortality": "descriptive",
    "t13_admission_vital_summary": "descriptive",
    "t03_severity_score_correlation": "association_or_qc",
    "t04_lactate_mortality_association": "association_or_qc",
    "t05_kdigo_renal_sensitivity": "association_or_qc",
    "t08_vaso_selection_bias_audit": "association_or_qc",
    "t09_sofa_zero_artefact_audit": "association_or_qc",
    "t10_complete_case_robustness": "association_or_qc",
    "t06_shock_phenotype_clustering": "advanced",
    "t07_mortality_prediction_auroc": "advanced",
    "t14_creatinine_trajectory_kdigo": "advanced",
    "t15_norepinephrine_dose_response": "advanced",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as fh:
        return [dict(row) for row in csv.DictReader(fh)]


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def write_text(path: Path, lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def task_num(task: str) -> int:
    m = re.match(r"t(\d+)_", task or "")
    return int(m.group(1)) if m else 999


def safe_git_commit(easyicu_root: Path) -> tuple[str, str]:
    try:
        full = subprocess.check_output(["git", "rev-parse", "HEAD"], cwd=easyicu_root, text=True).strip()
        return full, full[:7]
    except Exception:
        return "[unavailable]", "[unavailable]"


def extract_docx_paragraphs(path: Path) -> list[str]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    if not path.exists():
        return []
    with ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    out: list[str] = []
    for p in root.findall(".//w:p", ns):
        text = "".join((t.text or "") for t in p.findall(".//w:t", ns)).strip()
        if text:
            out.append(text)
    return out


def add_docx_paragraphs(doc: Document, markdown_lines: list[str]) -> None:
    for line in markdown_lines:
        if not line.strip():
            doc.add_paragraph("")
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            doc.add_paragraph(line)


def write_docx(path: Path, markdown_lines: list[str]) -> None:
    doc = Document()
    add_docx_paragraphs(doc, markdown_lines)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def summarize_sources(writing_root: Path, easyicu_root: Path, audit_dir: Path) -> dict[str, object]:
    matrix = read_csv(audit_dir / "matrix_status.csv")
    repair = read_csv(audit_dir / "repair_audit.csv")
    metrics = read_csv(audit_dir / "metric_sanity_audit.csv")
    task_summary = read_csv(audit_dir / "paper_task_metric_summary.csv")
    figure_source = read_csv(audit_dir / "curated_publication" / "source_maps" / "figure_source_map.csv")
    table_source = read_csv(audit_dir / "curated_publication" / "source_maps" / "table_source_map.csv")
    full_commit, short_commit = safe_git_commit(easyicu_root)
    return {
        "matrix": matrix,
        "repair": repair,
        "metrics": metrics,
        "task_summary": task_summary,
        "figure_source": figure_source,
        "table_source": table_source,
        "full_commit": full_commit,
        "short_commit": short_commit,
        "audit_dir": str(audit_dir),
        "writing_root": str(writing_root),
        "easyicu_root": str(easyicu_root),
    }


def make_table2(matrix: list[dict[str, str]], repair: list[dict[str, str]], task_summary: list[dict[str, str]]) -> list[dict[str, object]]:
    status = {(r.get("task_key"), r.get("arm")): r.get("status") for r in matrix}
    repair_count = Counter((r.get("task_key"), r.get("arm")) for r in repair)
    summary_by_task = {r.get("task_key"): r for r in task_summary}
    tasks = sorted({r.get("task_key") or "" for r in matrix}, key=task_num)
    rows: list[dict[str, object]] = []
    for task in tasks:
        row: dict[str, object] = {
            "task_key": task,
            "task_label": TASK_LABEL.get(task, task),
            "difficulty_band": DIFFICULTY.get(task, "other"),
            "primary_metric_summary": (summary_by_task.get(task) or {}).get("representative_metrics", "see audit table"),
            "total_repair_events": (summary_by_task.get(task) or {}).get("repair_events", ""),
        }
        for arm in ARM_ORDER:
            s = status.get((task, arm), "missing")
            rc = repair_count.get((task, arm), 0)
            row[f"{arm}_status"] = s
            row[f"{arm}_repairs"] = rc
        rows.append(row)
    return rows


def make_figure5_source(matrix: list[dict[str, str]], repair: list[dict[str, str]], metrics: list[dict[str, str]]) -> list[dict[str, object]]:
    repair_count = Counter((r.get("task_key"), r.get("arm")) for r in repair)
    metric_flags = Counter((r.get("task_key"), r.get("arm")) for r in metrics if str(r.get("manual_review_required") or "").lower() == "true")
    rows: list[dict[str, object]] = []
    for r in sorted(matrix, key=lambda x: (task_num(x.get("task_key") or ""), ARM_ORDER.index(x.get("arm") or "aware") if (x.get("arm") or "") in ARM_ORDER else 99)):
        task = r.get("task_key") or ""
        arm = r.get("arm") or ""
        rows.append({
            "task_key": task,
            "task_label": TASK_LABEL.get(task, task),
            "arm": arm,
            "arm_label": ARM_LABEL.get(arm, arm),
            "difficulty_band": DIFFICULTY.get(task, "other"),
            "acceptance_status": r.get("status") or "",
            "accepted": 1 if r.get("status") == "clean_ok" else 0,
            "repair_events": repair_count.get((task, arm), 0),
            "manual_metric_flags": metric_flags.get((task, arm), 0),
        })
    return rows


def make_tbd_map(summary: dict[str, object], out_dir: Path) -> list[dict[str, object]]:
    matrix = summary["matrix"]
    repair = summary["repair"]
    metrics = summary["metrics"]
    fig_src = summary["figure_source"]
    table_src = summary["table_source"]
    status_counts = Counter(r.get("status") for r in matrix)
    arm_clean = {arm: sum(1 for r in matrix if r.get("arm") == arm and r.get("status") == "clean_ok") for arm in ARM_ORDER}
    arm_total = {arm: sum(1 for r in matrix if r.get("arm") == arm) for arm in ARM_ORDER}
    manual_flags = sum(1 for r in metrics if str(r.get("manual_review_required") or "").lower() == "true")
    rows = [
        {"placeholder": "[TBD: full commit sha]", "replacement": summary["full_commit"], "source": "git rev-parse HEAD", "note": "current EASYICU repository commit at build time"},
        {"placeholder": "[TBD: full_commit_sha]", "replacement": summary["full_commit"], "source": "git rev-parse HEAD", "note": "same as above"},
        {"placeholder": "[TBD: batch_id]", "replacement": "v15_experiments_20260509_1854_full/final_audit_20260510_60clean", "source": str(summary["audit_dir"]), "note": "final audited v15 batch used for benchmark update"},
        {"placeholder": "[TBD: full_run_path]", "replacement": str(Path(str(summary["audit_dir"])).parent), "source": str(summary["audit_dir"]), "note": "preserved final run root"},
        {"placeholder": "[TBD: batch_date]", "replacement": "2026-05-10", "source": str(summary["audit_dir"]), "note": "final audit date"},
        {"placeholder": "[TBD: n_clean]", "replacement": status_counts.get("clean_ok", 0), "source": "matrix_status.csv", "note": "clean_ok final cells"},
        {"placeholder": "[TBD: n_clean] of 60", "replacement": f"{status_counts.get('clean_ok', 0)} of {len(matrix)}", "source": "matrix_status.csv", "note": "all task-arm cells"},
        {"placeholder": "[TBD: x/15] per-arm clean", "replacement": "; ".join(f"{arm}: {arm_clean[arm]}/{arm_total[arm]}" for arm in ARM_ORDER), "source": "matrix_status.csv", "note": "per-arm clean completion"},
        {"placeholder": "[TBD: repair/fallback events]", "replacement": len(repair), "source": "repair_audit.csv", "note": "repair burden, not model-autonomy success"},
        {"placeholder": "[TBD: metric sanity flags]", "replacement": manual_flags, "source": "metric_sanity_audit.csv", "note": "manual range review flags"},
        {"placeholder": "[TBD: figure records]", "replacement": len(fig_src), "source": "figure_source_map.csv", "note": "curated publication figure records"},
        {"placeholder": "[TBD: table records]", "replacement": len(table_src), "source": "table_source_map.csv", "note": "curated publication table records"},
        {"placeholder": "[TBD: Figure 5 output]", "replacement": "15-task × 4-arm benchmark panel package; source data in figure5_source_data.csv", "source": str(out_dir / "source_data" / "figure5_source_data.csv"), "note": "Figure 5 should be regenerated from this source"},
        {"placeholder": "[TBD: Figure 6 output]", "replacement": "two deep-dive cases: t04 clean association and t08 selection-bias audit", "source": str(out_dir / "figures" / "figure6_design.md"), "note": "Figure 6 should be regenerated from audited task evidence"},
    ]
    return rows


def main_manuscript_lines(summary: dict[str, object]) -> list[str]:
    matrix = summary["matrix"]
    repair = summary["repair"]
    metrics = summary["metrics"]
    fig_src = summary["figure_source"]
    table_src = summary["table_source"]
    clean = sum(1 for r in matrix if r.get("status") == "clean_ok")
    manual_flags = sum(1 for r in metrics if str(r.get("manual_review_required") or "").lower() == "true")
    arm_repairs = Counter(r.get("arm") for r in repair)
    lines = [
        "# EasyICU enables reproducible cross-database ICU phenotyping and evidence-bound agentic analysis",
        "",
        "Nature Methods v16 working draft generated from the user's EasyICU v15 Nature Methods manuscript, the EASYICU codebase, and the final v15 15-task audit.",
        "",
        "## Abstract",
        "",
        "Intensive-care-unit (ICU) datasets have become a major substrate for digital-medicine research, but analyses remain difficult to reproduce across sites because clinical concepts, time windows, aggregation rules and missingness semantics are encoded differently in each database. We developed EasyICU, an open-source toolkit that couples a cross-database ICU concept layer with preview-before-commit cohort construction, SOFA-2 rule translation, and an evidence-bound research-agent runtime. EasyICU normalises public ICU databases through a database abstraction and clinical concept layer; exposes interactive review of cohort definitions, features, trajectories and missingness; and exports analysis-ready cohorts with machine-readable metadata. A research-agent extension converts these exports into typed ResearchContext objects containing variable roles, units, time windows, allowed aggregations, missingness profiles and ICU-specific pitfalls. Model-generated analysis code is executed by a deterministic runtime, and every table, figure, statistic, script, log and manuscript claim is registered in a SHA-256 EvidenceStore before it can be cited. We demonstrate EasyICU in three complementary evaluations: a six-database source-cohort and SOFA-2 translation study, an evidence-bound SOFA-2 research-agent case study, and a v15 fifteen-task benchmark across four context/preference settings. In the final v15 audit, all 60 task-arm cells reached clean completion, with 230 recorded repair or fallback events, 60 metric-sanity checks and no manual range-review flags. These results support EasyICU as a reproducible infrastructure for ICU phenotyping and auditable agent-assisted analysis, while showing why final completion must be reported together with repair burden rather than as unassisted model performance.",
        "",
        "## 1 Introduction",
        "",
        "Public ICU databases make it possible to test clinical definitions and computational methods across large retrospective cohorts, but cross-database reproducibility is limited by heterogeneity in table structure, units, sampling frequency, missingness and clinical-rule implementation. A variable such as lactate, vasopressor exposure, urine output, creatinine or SOFA may look syntactically simple but can require database-specific source mappings, time-window definitions and aggregation rules before it is safe to compare across cohorts.",
        "",
        "EasyICU was built to make these steps explicit and reusable. Its first contribution is a clinical concept layer that maps source-specific ICU tables into standardised concepts and exposes those concepts through cohort-building and review workflows. Its second contribution is a preview-before-commit interface that lets users inspect source paths, cohort filters, extracted feature rows, patient-level time series, missingness and cross-database distributions before exporting a dataset. Its third contribution is a research-agent layer that turns exported cohorts and concept metadata into evidence-bound analyses rather than free-form model answers.",
        "",
        "The tool article therefore has a different emphasis from a conventional clinical-result manuscript. The primary object is not a new mortality model, a new lactate effect estimate or a norepinephrine dose-response curve. Those analyses are used as stress tests. The central claim is that EasyICU provides infrastructure for reproducible ICU phenotyping and for agent-assisted analyses whose code, outputs and manuscript statements remain traceable to registered evidence.",
        "",
        "## 2 System overview",
        "",
        "EasyICU has two connected layers. The platform layer supports cross-database ICU extraction, standardised concepts, cohort filtering, preview panels and exportable metadata. The agent layer consumes an EasyICU cohort and concept metadata to build a ResearchContext that drives planning, coding, validation, execution and evidence-bound writing. The codebase implements this separation explicitly: `easyicu.research_agent.context` builds ResearchContext objects from cohorts; `icu_rules.py` and related modules encode ICU variable roles, time windows and aggregation rules; `EvidenceStore` hashes and registers produced artefacts; and `ResearchAgentPipeline` orchestrates planning, execution, validation, writing, workflow graphs and deterministic replay bundles.",
        "",
        "Figure 1 should remain the architecture figure. Figure 2 should remain the preview-before-commit workflow figure. Figure 3 should remain the cross-database SOFA-2 component-activation figure. Figure 4 should remain the research-agent framework and evidence-bound reporting figure. Figures 5 and 6 are the only sections that should be updated using the final v15 benchmark.",
        "",
        "## 3 Results",
        "",
        "### 3.1 EasyICU unifies source cohorts across six public ICU databases",
        "",
        "The first results block should retain the existing v15 Nature Methods text describing the prepared source cohorts across MIMIC-IV, MIMIC-III, eICU-CRD, HiRID, AmsterdamUMCdb and SICdb. Table 1 remains the baseline-characteristics table for the full source cohorts before analytic filtering. This block establishes EasyICU as a cross-database platform rather than an analysis-only agent wrapper.",
        "",
        "### 3.2 Preview-before-commit makes extraction auditable before analysis",
        "",
        "The second results block should retain the existing description of EasyICU's extraction and review workflow. The key claim is that users can inspect data-source configuration, cohort filters, concept selection, feature-level preview rows, patient-level time series, denominator-aware missingness and cross-database distributions before committing an export. This is a methods/tool contribution and should remain upstream of any LLM benchmark.",
        "",
        "### 3.3 SOFA-2 demonstrates cross-database clinical-rule translation",
        "",
        "The third results block should retain the existing SOFA-2 case study. SOFA-2 is an appropriate stress test because it combines component-level rule translation, ordinal score handling, time-window choices and database-specific missingness. Figure 3 and the existing SOFA-1/SOFA-2 comparison table should continue to show that EasyICU can make rule translation visible across databases rather than silently collapsing everything into a single score.",
        "",
        "### 3.4 The research-agent layer converts EasyICU exports into evidence-bound analyses",
        "",
        "The fourth results block should describe the research-agent layer as an extension of the EasyICU platform. The agent receives structured metadata rather than raw rows in the prompt. It can plan and write code, but computed values come from executed Python scripts. The pipeline registers results_report.md, manifest.json, workflow_graph.json, execution_replay.json, audit logs, tables, figures, statistics and manuscript scaffold files. The manuscript writer is constrained to cite registered evidence ids, and clinical interpretation remains human supervised.",
        "",
        "### 3.5 A fifteen-task benchmark stress-tests the evidence-bound agent layer",
        "",
        f"The final v15 benchmark should replace the older ten-task or six-task ablation language. It evaluated 15 real-cohort ICU analysis tasks under four settings: ICU-aware with preferences, ICU-aware without preferences, generic context with preferences and generic context without preferences. The design produced {len(matrix)} task-arm cells. In the final audit, {clean} of {len(matrix)} cells reached clean completion. Automated metric sanity checks covered {len(metrics)} cells and flagged {manual_flags} cells for manual range review. The final audit recorded {len(repair)} repair or fallback events. Curated publication outputs included {len(fig_src)} figure records and {len(table_src)} table records with source maps to the original run evidence.",
        "",
        "Figure 5 should present this benchmark as a tool-validation panel, not as the main scientific discovery. Panel a should show the 15 × 4 clean-completion heatmap. Panel b should show vigilance and repair burden by arm. Panel c should summarise completion and repair burden across descriptive, association/QC and advanced-analysis task bands. Panel d should show the two-factor context/preference decomposition. Because all final cells reached clean completion, Figure 5 should emphasise repair burden, auditability and system assistance rather than final success alone.",
        "",
        "Repair/fallback events by setting were:",
        "",
    ]
    for arm in ARM_ORDER:
        lines.append(f"- {ARM_LABEL[arm]} (`{arm}`): {arm_repairs.get(arm, 0)} repair/fallback events.")
    lines.extend([
        "",
        "### 3.6 Deep-dive cases show how EasyICU separates clean output from audit interpretation",
        "",
        "Figure 6 should use two representative v15 tasks to show the audit trail rather than to claim new clinical effects. The left column should use t04 lactate-mortality association as a clean association case, including an adjusted death-probability curve, an odds-ratio panel and an evidence-provenance panel. The right column should use t08 vasopressor selection-bias audit as a divergent case, including per-arm vasopressor association estimates, warning or repair timelines and per-arm acceptance status. This design highlights the value of EasyICU's evidence-bound runtime: the system does not simply return a plot, but also reports where the analysis came from and where interpretation should be cautious.",
        "",
        "## 4 Discussion",
        "",
        "EasyICU addresses a gap between ICU database preprocessing tools and general-purpose agentic analysis systems. Existing preprocessing workflows can standardise data, and general LLM agents can produce analysis code, but high-stakes clinical research needs the two pieces to be coupled through explicit clinical semantics and provenance. EasyICU contributes that coupling: the same concept metadata used to extract and review cohorts is passed forward into the research-agent context, validators and evidence-bound manuscript scaffold.",
        "",
        "The v15 benchmark should be interpreted as validation of this infrastructure. The 60/60 clean-completion result demonstrates that the final audited workflow can produce complete outputs across heterogeneous ICU tasks. The 230 recorded repair/fallback events are equally important, because they show that an execution-audited system should report how much deterministic assistance was required. This is more transparent than describing final outputs as unassisted LLM performance.",
        "",
        "EasyICU also makes clinical-rule ambiguity visible. The SOFA-2 case study and SOFA-zero audit illustrate that low numeric scores can encode missing components rather than low illness severity. The research-agent layer is designed around this principle: agents receive explicit aggregation rules and known pitfalls, validators check generated scripts, and manuscript claims are bound to evidence rather than left as unsupported prose.",
        "",
        "Several limitations remain. EasyICU currently depends on local database preparation and source mapping quality. The research-agent layer should not be described as a fully autonomous clinical scientist; it is an auditable assistant whose outputs require human clinical and statistical review. The v15 benchmark was run on a single local open-weight coding model and should be expanded across models and external sites before being presented as a public leaderboard. Finally, association tasks in the benchmark are not causal analyses and should not be interpreted as treatment effects.",
        "",
        "## 5 Methods overview for the v16 update",
        "",
        "The unchanged Methods sections should continue to describe database preparation, concept extraction, cohort filtering, SOFA-2 rule implementation, preview workflow and research-agent runtime. The v16 update should add a concise fifteen-task protocol subsection: task registry, four-arm context/preference design, clean-completion definition, metric sanity audit, repair/fallback taxonomy, figure/table curation and evidence source mapping.",
        "",
        "## Data availability",
        "",
        "Use the existing v15 Nature Methods data-availability language for public ICU databases and local prepared exports. Add that the final v15 benchmark audit and derived manuscript evidence package are stored under the preserved final run directory listed in the provenance files.",
        "",
        "## Code availability",
        "",
        f"Use the existing code-availability language and update the canonical EASYICU commit placeholder to `{summary['full_commit']}` when the author confirms this build commit is the submission reference.",
        "",
        "## References",
        "",
        "Keep the references from the existing v15 Word manuscript. Do not import or fabricate new references automatically without a separate literature-check step.",
    ])
    return lines


def supplement_lines(summary: dict[str, object]) -> list[str]:
    matrix = summary["matrix"]
    repair = summary["repair"]
    metrics = summary["metrics"]
    phases = Counter(r.get("phase") or "unknown" for r in repair)
    task_repairs = Counter(r.get("task_key") for r in repair)
    lines = [
        "# EasyICU Nature Methods v16 supplementary update",
        "",
        "This file is not a replacement for the full v15 supplementary Word document. It is the v16 update block that should be merged into the existing supplementary materials while preserving S1-S9, S5b, M1 and M2 unless manually superseded.",
        "",
        "## Supplementary Methods M3. Fifteen-task four-arm research-agent benchmark",
        "",
        "The v15 benchmark evaluated the EasyICU research-agent layer on 15 real-cohort ICU analysis tasks. The tasks spanned descriptive summaries, stratified outcome summaries, severity-score correlation, association modelling, renal sensitivity analysis, phenotype clustering, prediction performance, selection-bias audit, SOFA-zero artefact audit, missing-data robustness, length-of-stay summaries, age-stratified mortality, vital-sign summaries, creatinine trajectory/KDIGO analysis and norepinephrine-equivalent dose-response analysis.",
        "",
        "Each task was evaluated under four settings: ICU-aware with user preferences, ICU-aware without user preferences, generic context with user preferences and generic context without user preferences. The ICU-aware settings supplied EasyICU ResearchContext metadata including variable roles, units, time windows, missingness profiles, allowed aggregations and known ICU pitfalls. The generic settings retained only a reduced context analogous to a general data-agent prompt. The preference factor tested whether explicit user preferences changed completion or repair burden.",
        "",
        "A task-arm cell reached clean completion when execution finished, expected artifacts were present, key metrics were extractable, and no unresolved contract-level failure remained. Repair/fallback events were recorded separately and should be interpreted as system assistance required to obtain the final audited output.",
        "",
        "## Supplementary Results update",
        "",
        f"The final audit contained {len(matrix)} task-arm cells. All {len(matrix)} reached clean completion. Metric sanity checks covered {len(metrics)} cells. The final audit recorded {len(repair)} repair or fallback events.",
        "",
        "### Repair/fallback events by phase",
        "",
    ]
    for phase, count in phases.most_common():
        lines.append(f"- {phase}: {count}")
    lines.extend([
        "",
        "### Repair/fallback events by task",
        "",
    ])
    for task, count in sorted(task_repairs.items(), key=lambda kv: task_num(kv[0] or "")):
        lines.append(f"- {task}: {count}")
    lines.extend([
        "",
        "## Supplementary Table S10. Per-run acceptance summary",
        "",
        "Use `tables/S10_per_run_acceptance_summary.csv`, generated from `matrix_status.csv` with per-cell repair counts and metric-review flags.",
        "",
        "## Supplementary Table S11. Per-run context-ablation audit",
        "",
        "Use `tables/S11_context_ablation_audit.csv`, which adds context/preference factors and task difficulty bands for the 15 × 4 design.",
        "",
        "## Supplementary Table S12. Fifteen-task numerical-results matrix",
        "",
        "Use `tables/S12_fifteen_task_numerical_results_matrix.csv`, generated from the final paper task metric summary. This table is for audit and supplement, not for overinterpreting task-specific clinical associations.",
    ])
    return lines


def design_files(summary: dict[str, object], out_dir: Path) -> None:
    matrix = summary["matrix"]
    repair = summary["repair"]
    metrics = summary["metrics"]
    clean = sum(1 for r in matrix if r.get("status") == "clean_ok")
    manual_flags = sum(1 for r in metrics if str(r.get("manual_review_required") or "").lower() == "true")
    write_text(out_dir / "02_v16_restructure_plan.md", [
        "# v16 Nature Methods restructure plan",
        "",
        "## Correct manuscript identity",
        "",
        "This is an EasyICU tool/methods paper, not a standalone v15 experiment paper.",
        "",
        "## Main claim",
        "",
        "EasyICU enables reproducible cross-database ICU phenotyping and evidence-bound agentic analysis by connecting concept-level extraction, preview-before-commit review, SOFA-2 rule translation, ResearchContext generation, deterministic execution, validators and SHA-256 evidence binding.",
        "",
        "## What to preserve from v15 Word",
        "",
        "- Introduction and tool motivation.",
        "- Six-database source cohort and Table 1.",
        "- Preview-before-commit workflow and Figure 2.",
        "- SOFA-2 cross-database rule translation and Figure 3.",
        "- Research-agent framework and Figure 4.",
        "- Existing references unless separately checked.",
        "",
        "## What to replace",
        "",
        "- Replace old six-task or ten-task benchmark language with the final 15-task × 4-arm v15 audit.",
        "- Replace outdated Figure 5 / Figure 6 captions with the new benchmark and deep-dive case designs.",
        "- Replace Table 2 benchmark body with the v15 15×4 table design.",
        "- Add Supplementary Methods M3 and Supplementary Tables S10-S12.",
    ])
    write_text(out_dir / "figures" / "figure5_design.md", [
        "# Figure 5 design — final v15 15-task benchmark",
        "",
        f"Source audit: {summary['audit_dir']}",
        "",
        "## Core claim",
        "",
        f"The EasyICU research-agent layer produced complete, auditable outputs across {len(matrix)} task-arm cells, but interpretation must include repair burden ({len(repair)} repair/fallback events).",
        "",
        "## Panels",
        "",
        "- **a. Acceptance heatmap**: 15 tasks × 4 arms, colored by final status. All 60 cells should be clean_ok in the final audited run.",
        "- **b. Vigilance scorecard**: repair/fallback events by arm, plus metric manual-review flags.",
        "- **c. Difficulty-stratified acceptance**: descriptive, association/QC and advanced task bands by arm.",
        "- **d. 2×2 factor decomposition**: ICU-aware versus generic context, with versus without preferences; plot mean repair burden and clean completion.",
        "",
        "## Numbers to show",
        "",
        f"- clean_ok: {clean}/{len(matrix)}",
        f"- metric sanity cells: {len(metrics)}",
        f"- manual metric flags: {manual_flags}",
        f"- repair/fallback events: {len(repair)}",
    ])
    write_text(out_dir / "figures" / "figure6_design.md", [
        "# Figure 6 design — two deep-dive agentic analysis cases",
        "",
        "## Core claim",
        "",
        "EasyICU makes agent outputs traceable and audit-ready: a clean association case and a divergent selection-bias case can both be inspected through evidence, warnings, repairs and provenance.",
        "",
        "## Left column: t04 lactate-mortality association",
        "",
        "- **a. Adjusted death probability versus lactate_max_24h**.",
        "- **b. Odds-ratio forest plot across relevant arms or model specifications**.",
        "- **c. Evidence provenance panel showing source step summary, figure source map and registered artifacts**.",
        "",
        "## Right column: t08 vasopressor selection-bias audit",
        "",
        "- **d. Adjusted vasopressor or mortality association estimates by arm or stratum**.",
        "- **e. Warning / repair timeline or scorecard by arm**.",
        "- **f. Per-arm acceptance status and audit note**.",
        "",
        "## Wording constraint",
        "",
        "These are tool-validation cases, not new causal clinical claims.",
    ])
    write_text(out_dir / "tables" / "table2_design.md", [
        "# Table 2 design — 15-task × 4-arm benchmark summary",
        "",
        "## Columns",
        "",
        "- task key and task label",
        "- difficulty band",
        "- primary metric summary from final audit",
        "- per-arm clean status",
        "- per-arm repair/fallback event count",
        "- total repair/fallback event count",
        "",
        "## Interpretation",
        "",
        "Because all final cells are clean_ok, Table 2 should not be read as a simple success/failure comparison. Its value is showing repair burden and audit traceability across task types and context/preference settings.",
    ])
    write_text(out_dir / "07_supplement_update_plan.md", [
        "# Supplement update plan",
        "",
        "## Preserve",
        "",
        "- Supplementary Figures S1-S2 unless manually superseded.",
        "- Supplementary Tables S1-S9 and S5b unless manually superseded.",
        "- Supplementary Methods M1 runtime and M2 research-agent details.",
        "",
        "## Add or replace",
        "",
        "- M3 fifteen-task four-arm benchmark protocol.",
        "- S10 per-run acceptance summary, 60 rows.",
        "- S11 context-ablation audit, 60 rows.",
        "- S12 fifteen-task numerical-results matrix.",
    ])


def write_extracted_outline(writing_root: Path, out_dir: Path) -> None:
    files = [
        writing_root / "EasyICU_main_manuscript_NatureMethods_v15_20260510.docx",
        writing_root / "EasyICU_supplementary_materials_NatureMethods_v15_20260510.docx",
        writing_root / "EasyICU_main_manuscript_npjDM_v14_real_cohort_10task_20260509.docx",
        writing_root / "_archived_pre_v14_20260509" / "EasyICU_main_manuscript_npjDM_v9_with_agent_20260505.docx",
    ]
    lines = ["# Extracted source manuscript outline", ""]
    for path in files:
        paras = extract_docx_paragraphs(path)
        lines.extend([f"## {path.name}", "", f"Path: `{path}`", "", f"Paragraphs extracted: {len(paras)}", ""])
        for text in paras:
            if re.match(r"^(Abstract|Introduction|Results|Discussion|Methods|Data availability|Code availability|Author contributions|References)$", text, re.I) or re.match(r"^(Figure|Table) \d+\.", text) or "[TBD:" in text:
                lines.append(f"- {text[:900]}")
        lines.append("")
    write_text(out_dir / "01_extracted_current_nm_v15.md", lines)


def build(args: argparse.Namespace) -> None:
    writing_root = Path(args.writing_root).resolve()
    easyicu_root = Path(args.easyicu_root).resolve()
    audit_dir = Path(args.audit_dir).resolve()
    out_dir = Path(args.out_dir).resolve() if args.out_dir else writing_root / "_nature_methods_v16_restructure_20260511"
    out_dir.mkdir(parents=True, exist_ok=True)
    summary = summarize_sources(writing_root, easyicu_root, audit_dir)
    table2 = make_table2(summary["matrix"], summary["repair"], summary["task_summary"])
    table2_fields = ["task_key", "task_label", "difficulty_band", "primary_metric_summary", "total_repair_events"]
    for arm in ARM_ORDER:
        table2_fields.extend([f"{arm}_status", f"{arm}_repairs"])
    write_csv(out_dir / "tables" / "Table2_15x4_benchmark_v16.csv", table2, table2_fields)
    fig5_source = make_figure5_source(summary["matrix"], summary["repair"], summary["metrics"])
    write_csv(out_dir / "source_data" / "figure5_source_data.csv", fig5_source, ["task_key", "task_label", "arm", "arm_label", "difficulty_band", "acceptance_status", "accepted", "repair_events", "manual_metric_flags"])
    s10 = fig5_source
    write_csv(out_dir / "tables" / "S10_per_run_acceptance_summary.csv", s10, ["task_key", "task_label", "arm", "arm_label", "difficulty_band", "acceptance_status", "accepted", "repair_events", "manual_metric_flags"])
    s11 = []
    for row in fig5_source:
        arm = str(row["arm"])
        s11.append({**row, "icu_context": 1 if arm in {"aware", "aware_no_pref"} else 0, "user_preferences": 1 if arm in {"aware", "naive_with_pref"} else 0})
    write_csv(out_dir / "tables" / "S11_context_ablation_audit.csv", s11, ["task_key", "task_label", "arm", "arm_label", "icu_context", "user_preferences", "difficulty_band", "acceptance_status", "accepted", "repair_events", "manual_metric_flags"])
    write_csv(out_dir / "tables" / "S12_fifteen_task_numerical_results_matrix.csv", summary["task_summary"], list(summary["task_summary"][0].keys()) if summary["task_summary"] else ["empty"])
    tbd_rows = make_tbd_map(summary, out_dir)
    write_csv(out_dir / "03_tbd_substitution_map.csv", tbd_rows, ["placeholder", "replacement", "source", "note"])
    write_extracted_outline(writing_root, out_dir)
    design_files(summary, out_dir)
    main_lines = main_manuscript_lines(summary)
    supp_lines = supplement_lines(summary)
    write_text(out_dir / "EasyICU_main_manuscript_NatureMethods_v16_20260511.md", main_lines)
    write_docx(out_dir / "EasyICU_main_manuscript_NatureMethods_v16_20260511.docx", main_lines)
    write_text(out_dir / "EasyICU_supplementary_materials_NatureMethods_v16_20260511_update.md", supp_lines)
    write_docx(out_dir / "EasyICU_supplementary_materials_NatureMethods_v16_20260511_update.docx", supp_lines)
    readme = [
        "# EasyICU Nature Methods v16 restructure package",
        "",
        "This package corrects the manuscript direction: EasyICU is a tool/methods article, while the final v15 benchmark is one validation component.",
        "",
        "## Start here",
        "",
        "1. `EasyICU_main_manuscript_NatureMethods_v16_20260511.md`",
        "2. `EasyICU_main_manuscript_NatureMethods_v16_20260511.docx`",
        "3. `03_tbd_substitution_map.csv`",
        "4. `figures/figure5_design.md`",
        "5. `figures/figure6_design.md`",
        "6. `tables/Table2_15x4_benchmark_v16.csv`",
        "7. `EasyICU_supplementary_materials_NatureMethods_v16_20260511_update.md`",
        "",
        "## Important",
        "",
        "The generated v16 draft is an author-review draft. It intentionally preserves the EasyICU tool-paper framing and does not turn v15 into a standalone clinical-results manuscript.",
    ]
    write_text(out_dir / "00_README_START_HERE.md", readme)
    payload = {
        "generated_at": datetime.now().isoformat(),
        "out_dir": str(out_dir),
        "matrix_rows": len(summary["matrix"]),
        "clean_ok": sum(1 for r in summary["matrix"] if r.get("status") == "clean_ok"),
        "repair_events": len(summary["repair"]),
        "metric_rows": len(summary["metrics"]),
        "figure_source_rows": len(summary["figure_source"]),
        "table_source_rows": len(summary["table_source"]),
        "easyicu_commit": summary["full_commit"],
    }
    (out_dir / "build_summary.json").write_text(json.dumps(payload, indent=2, ensure_ascii=False) + "\n", encoding="utf-8")
    print(out_dir)
    print(out_dir / "00_README_START_HERE.md")
    print(out_dir / "EasyICU_main_manuscript_NatureMethods_v16_20260511.docx")
    print(out_dir / "build_summary.json")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--writing-root", default="/Users/haibo/Documents/GitHub/easyicu写作/00_当前文章/01_当前正式稿")
    parser.add_argument("--easyicu-root", default="/Users/haibo/Documents/GitHub/EASYICU")
    parser.add_argument("--audit-dir", default="/Users/haibo/Documents/GitHub/EASYICU/research_output/v15_experiments_20260509_1854_full/final_audit_20260510_60clean")
    parser.add_argument("--out-dir", default=None)
    build(parser.parse_args())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
