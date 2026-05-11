#!/usr/bin/env python3
from __future__ import annotations

import argparse
import csv
import json
import re
from collections import Counter, defaultdict
from datetime import datetime
from pathlib import Path
from zipfile import ZipFile
import xml.etree.ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ARM_ORDER = ["aware", "aware_no_pref", "naive_with_pref", "naive"]
ARM_LABEL = {
    "aware": "ICU-aware + preferences",
    "aware_no_pref": "ICU-aware only",
    "naive_with_pref": "Generic context + preferences",
    "naive": "Generic context only",
}
TASK_LABEL = {
    "t01_table_one_descriptive": "baseline descriptive summary",
    "t02_outcome_incidence_strata": "outcome incidence by severity strata",
    "t03_severity_score_correlation": "severity-score correlation",
    "t04_lactate_mortality_association": "lactate-mortality association",
    "t05_kdigo_renal_sensitivity": "KDIGO renal sensitivity",
    "t06_shock_phenotype_clustering": "shock phenotype clustering",
    "t07_mortality_prediction_auroc": "mortality prediction performance",
    "t08_vaso_selection_bias_audit": "vasopressor selection-bias audit",
    "t09_sofa_zero_artefact_audit": "SOFA-zero artefact audit",
    "t10_complete_case_robustness": "complete-case robustness",
    "t11_los_distribution_descriptive": "ICU length-of-stay distribution",
    "t12_age_stratified_mortality": "age-stratified mortality",
    "t13_admission_vital_summary": "admission vital-summary profile",
    "t14_creatinine_trajectory_kdigo": "creatinine trajectory and KDIGO analysis",
    "t15_norepinephrine_dose_response": "norepinephrine-equivalent dose-response analysis",
}
TASK_BAND = {
    "t01_table_one_descriptive": "descriptive",
    "t02_outcome_incidence_strata": "descriptive",
    "t11_los_distribution_descriptive": "descriptive",
    "t12_age_stratified_mortality": "descriptive",
    "t13_admission_vital_summary": "descriptive",
    "t03_severity_score_correlation": "association or quality-control",
    "t04_lactate_mortality_association": "association or quality-control",
    "t05_kdigo_renal_sensitivity": "association or quality-control",
    "t08_vaso_selection_bias_audit": "association or quality-control",
    "t09_sofa_zero_artefact_audit": "association or quality-control",
    "t10_complete_case_robustness": "association or quality-control",
    "t06_shock_phenotype_clustering": "advanced analysis",
    "t07_mortality_prediction_auroc": "advanced analysis",
    "t14_creatinine_trajectory_kdigo": "advanced analysis",
    "t15_norepinephrine_dose_response": "advanced analysis",
}


def read_csv(path: Path) -> list[dict[str, str]]:
    if not path.exists():
        return []
    with path.open(newline="", encoding="utf-8") as fh:
        return [dict(row) for row in csv.DictReader(fh)]


def write_csv(path: Path, rows: list[dict[str, object]], fields: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", newline="", encoding="utf-8") as fh:
        writer = csv.DictWriter(fh, fieldnames=fields, extrasaction="ignore")
        writer.writeheader()
        writer.writerows(rows)


def write_text(path: Path, lines: list[str]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("\n".join(lines).rstrip() + "\n", encoding="utf-8")


def task_num(task: str) -> int:
    m = re.match(r"t(\d+)_", task or "")
    return int(m.group(1)) if m else 999


def extract_docx_text(path: Path) -> list[str]:
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    if not path.exists():
        return []
    with ZipFile(path) as zf:
        root = ET.fromstring(zf.read("word/document.xml"))
    paragraphs: list[str] = []
    for p in root.findall(".//w:p", ns):
        text = "".join((t.text or "") for t in p.findall(".//w:t", ns)).strip()
        if text:
            paragraphs.append(text)
    return paragraphs


def markdown_to_docx(lines: list[str], path: Path) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    styles = doc.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(10.5)
    for line in lines:
        raw = line.rstrip()
        if not raw:
            doc.add_paragraph("")
            continue
        if raw.startswith("# "):
            p = doc.add_heading(raw[2:].strip(), level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif raw.startswith("## "):
            doc.add_heading(raw[3:].strip(), level=1)
        elif raw.startswith("### "):
            doc.add_heading(raw[4:].strip(), level=2)
        elif raw.startswith("#### "):
            doc.add_heading(raw[5:].strip(), level=3)
        elif raw.startswith("- "):
            doc.add_paragraph(raw[2:].strip(), style="List Bullet")
        elif raw.startswith("| "):
            doc.add_paragraph(raw)
        else:
            doc.add_paragraph(raw)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(path)


def status_table(matrix: list[dict[str, str]], repair: list[dict[str, str]]) -> list[dict[str, object]]:
    repairs = Counter((r.get("task_key"), r.get("arm")) for r in repair)
    status = {(r.get("task_key"), r.get("arm")): r.get("status") for r in matrix}
    tasks = sorted({r.get("task_key") or "" for r in matrix}, key=task_num)
    rows: list[dict[str, object]] = []
    for task in tasks:
        row: dict[str, object] = {
            "Task": task,
            "Plain-language task": TASK_LABEL.get(task, task),
            "Task band": TASK_BAND.get(task, "other"),
        }
        for arm in ARM_ORDER:
            row[f"{arm} status"] = status.get((task, arm), "missing")
            row[f"{arm} repairs"] = repairs.get((task, arm), 0)
        row["Total repairs"] = sum(repairs.get((task, arm), 0) for arm in ARM_ORDER)
        rows.append(row)
    return rows


def csv_markdown_table(rows: list[dict[str, object]], fields: list[str], max_rows: int | None = None) -> list[str]:
    if max_rows is not None:
        rows = rows[:max_rows]
    out = ["| " + " | ".join(fields) + " |", "| " + " | ".join("---" for _ in fields) + " |"]
    for row in rows:
        vals = []
        for f in fields:
            value = str(row.get(f, ""))
            value = value.replace("|", "/").replace("\n", " ")
            vals.append(value)
        out.append("| " + " | ".join(vals) + " |")
    return out


def load_context(args: argparse.Namespace) -> dict[str, object]:
    base = Path(args.v16_dir).resolve()
    audit = Path(args.audit_dir).resolve()
    writing = Path(args.writing_root).resolve()
    matrix = read_csv(audit / "matrix_status.csv")
    repair = read_csv(audit / "repair_audit.csv")
    metrics = read_csv(audit / "metric_sanity_audit.csv")
    task_summary = read_csv(audit / "paper_task_metric_summary.csv")
    figure_source = read_csv(audit / "curated_publication" / "source_maps" / "figure_source_map.csv")
    table_source = read_csv(audit / "curated_publication" / "source_maps" / "table_source_map.csv")
    tbd = read_csv(base / "03_tbd_substitution_map.csv")
    current_main = (base / "EasyICU_main_manuscript_NatureMethods_v16_20260511.md").read_text(encoding="utf-8") if (base / "EasyICU_main_manuscript_NatureMethods_v16_20260511.md").exists() else ""
    current_supp = (base / "EasyICU_supplementary_materials_NatureMethods_v16_20260511_update.md").read_text(encoding="utf-8") if (base / "EasyICU_supplementary_materials_NatureMethods_v16_20260511_update.md").exists() else ""
    v15_main = extract_docx_text(writing / "EasyICU_main_manuscript_NatureMethods_v15_20260510.docx")
    v15_supp = extract_docx_text(writing / "EasyICU_supplementary_materials_NatureMethods_v15_20260510.docx")
    return {
        "base": base,
        "audit": audit,
        "writing": writing,
        "matrix": matrix,
        "repair": repair,
        "metrics": metrics,
        "task_summary": task_summary,
        "figure_source": figure_source,
        "table_source": table_source,
        "tbd": tbd,
        "current_main": current_main,
        "current_supp": current_supp,
        "v15_main": v15_main,
        "v15_supp": v15_supp,
    }


def build_main_lines(ctx: dict[str, object]) -> list[str]:
    matrix = ctx["matrix"]
    repair = ctx["repair"]
    metrics = ctx["metrics"]
    figure_source = ctx["figure_source"]
    table_source = ctx["table_source"]
    clean = sum(1 for r in matrix if r.get("status") == "clean_ok")
    manual_flags = sum(1 for r in metrics if str(r.get("manual_review_required") or "").lower() == "true")
    repairs_by_arm = Counter(r.get("arm") for r in repair)
    repairs_by_task = Counter(r.get("task_key") for r in repair)
    table_rows = status_table(matrix, repair)
    title = "EasyICU enables reproducible cross-database ICU phenotyping and evidence-bound agentic analysis"
    lines = [
        f"# {title}",
        "",
        "Authors: [author list to be confirmed]",
        "",
        "Article type: Nature Methods tool/methods article",
        "",
        "## Abstract",
        "",
        "Intensive-care-unit (ICU) databases are widely used for digital-medicine research, but analyses remain difficult to reproduce because clinical concepts, time windows, aggregation rules, units and missingness semantics differ across databases. We developed EasyICU, an open-source toolkit for cross-database ICU phenotyping, preview-before-commit cohort construction, SOFA-2 rule translation and evidence-bound agentic analysis. EasyICU separates the platform layer, which maps ICU data sources into reviewable clinical concepts and exportable cohorts, from the research-agent layer, which converts exported cohorts into structured ResearchContext objects, executes model-generated analysis code in a deterministic runtime, and registers every table, figure, statistic, script, log and manuscript-relevant claim in a SHA-256 EvidenceStore. We evaluate EasyICU through six-database source-cohort preparation, a SOFA-2 rule-translation case study, an evidence-bound SOFA-2 research-agent demonstration and a final fifteen-task benchmark across four context/preference settings. In the final audit, all 60 task-arm cells reached clean completion; 230 repair or fallback events were recorded; 60 metric-sanity checks produced no manual range-review flags; and curated publication outputs included 16 figure records and seven table records with source maps. These results support EasyICU as an infrastructure for reproducible ICU cohort construction and auditable agent-assisted analysis. They also show why final completion should be reported together with repair burden rather than described as unassisted model performance.",
        "",
        "## Introduction",
        "",
        "Critical-care research increasingly depends on large retrospective ICU databases. These resources make it possible to evaluate severity scores, organ-dysfunction definitions, treatment-pattern summaries and predictive models across large numbers of ICU stays. However, the practical work needed to make these analyses reproducible remains difficult. The same clinical concept can be stored under different item identifiers, units, sampling frequencies and table structures. Time windows must be chosen before aggregation. Missingness can reflect measurement practice rather than absence of disease. Rules such as SOFA and KDIGO depend on multiple components and clinically meaningful defaults. If these choices are hidden inside one-off scripts, an apparently simple cohort can become hard to inspect, reproduce or translate across databases.",
        "",
        "EasyICU addresses this problem by treating ICU phenotyping as a tool and workflow problem rather than only a statistical-modelling problem. The platform maps database-specific tables into standardised clinical concepts, exposes those concepts through cohort-building and preview interfaces, and exports analysis-ready cohorts with machine-readable metadata. A user can review source paths, cohort filters, feature previews, patient-level trajectories, missingness summaries and cross-database distributions before committing to an analysis. This design is intended to make clinical-rule implementation visible before modelling begins.",
        "",
        "The second part of EasyICU links these reviewed exports to an evidence-bound research-agent runtime. General-purpose code-writing agents can produce useful analysis scripts, but clinical research requires stronger guardrails: the agent should know which variables are outcomes, exposures, confounders or quality-control fields; which units and time windows are valid; which aggregations are clinically sensible; and which claims can be supported by executed evidence. EasyICU therefore packages exported cohorts and concept metadata into ResearchContext objects, validates generated analysis plans and scripts against ICU-specific rules, executes them under a runtime supervisor and binds manuscript-facing claims to registered evidence artifacts.",
        "",
        "The central contribution of this manuscript is EasyICU as a method and tool. The v15 benchmark is used as a stress test of the research-agent layer, not as the primary scientific discovery. Individual lactate, vasopressor, creatinine or norepinephrine results are included only to demonstrate the workflow and audit trail. We therefore report final completion, metric sanity and repair burden together, and we avoid interpreting association analyses as causal treatment effects.",
        "",
        "## Results",
        "",
        "### EasyICU source cohorts provide a reproducible cross-database ICU foundation",
        "",
        "EasyICU first standardises the process of preparing source cohorts from public ICU databases. The manuscript should retain the detailed v15 cohort descriptions for MIMIC-IV, MIMIC-III, eICU-CRD, HiRID, AmsterdamUMCdb and SICdb. In the v16 framing, this section is the foundation of the paper: it shows that EasyICU is not simply an LLM wrapper, but a data and concept infrastructure for ICU research. Table 1 should report the six-database source-cohort characteristics before task-specific filtering, including the cohort denominators, ICU-stay or patient counts, outcome availability and clinically important missingness summaries already present in the v15 draft.",
        "",
        "### Preview-before-commit review makes cohort construction auditable",
        "",
        "The platform exposes a preview-before-commit workflow so that users can inspect concept extraction before using an exported cohort. Figure 2 should show the workflow from database configuration to concept selection, feature preview, trajectory review, missingness review, cross-database comparison and export. The most important message is simple: EasyICU lets the researcher look at what will be analysed before it becomes a fixed dataset. This supports reproducibility because cohort definitions, aggregation choices and export metadata are made visible at the point where mistakes are easiest to correct.",
        "",
        "### SOFA-2 illustrates cross-database clinical-rule translation",
        "",
        "SOFA-2 is retained as the main clinical-rule case study. It is suitable for a Nature Methods tool paper because it combines several challenges that are common in ICU research: component-level rule translation, multiple physiologic and laboratory inputs, ordinal scoring, time-window selection and missingness-sensitive interpretation. Figure 3 should show component activation and cross-database availability. The text should emphasize that EasyICU does not only compute a score; it exposes which components are available, which rules are activated and where a low score may reflect missing source components rather than truly low severity.",
        "",
        "### The research-agent layer converts EasyICU exports into evidence-bound analyses",
        "",
        "The research-agent layer extends EasyICU from cohort construction to auditable analysis. It receives structured metadata rather than only raw table columns. The ResearchContext contains variable roles, units, time windows, known pitfalls, missingness profiles and allowed aggregation rules. The planner proposes a typed analysis plan; the coder writes scripts; the runtime supervisor executes them; validators inspect outputs; and the writer can only support result-like claims through registered evidence. The EvidenceStore records SHA-256 hashes and provenance for scripts, logs, figures, tables and manuscript-scaffold files. Figure 4 should show this sequence as a tool workflow rather than as a black-box autonomous agent.",
        "",
        "### Fifteen-task benchmark stress-tests the research-agent layer",
        "",
        f"The final v15 benchmark evaluated 15 real-cohort ICU analysis tasks under four settings: ICU-aware with preferences, ICU-aware without preferences, generic context with preferences and generic context without preferences. This yielded {len(matrix)} task-arm cells. In the final audit, {clean} of {len(matrix)} cells reached clean completion. Metric sanity checks covered {len(metrics)} cells and produced {manual_flags} manual range-review flags. The audit recorded {len(repair)} repair or fallback events. Curated publication outputs included {len(figure_source)} figure records and {len(table_source)} table records, each linked to source maps in the final audit directory.",
        "",
        "The benchmark should be interpreted as a tool-validation experiment. All final cells completed cleanly, but this result alone would hide the amount of deterministic repair and fallback needed to obtain a clean final package. Figure 5 should therefore show both acceptance and vigilance: the 15 × 4 completion matrix, repair burden by arm, task-band summaries and the two-factor context/preference design. Table 2 reports per-task status and repair counts.",
        "",
        "Repair/fallback events by setting were:",
        "",
    ]
    for arm in ARM_ORDER:
        lines.append(f"- {ARM_LABEL[arm]}: {repairs_by_arm.get(arm, 0)} events.")
    lines.extend([
        "",
        "The highest repair burdens were concentrated in tasks that required more brittle plotting, missing-data handling, trajectory analysis or dependency-sensitive modelling:",
        "",
    ])
    for task, count in sorted(repairs_by_task.items(), key=lambda kv: (-kv[1], task_num(kv[0] or "")))[:6]:
        lines.append(f"- {task}: {count} events ({TASK_LABEL.get(task, task)}).")
    lines.extend([
        "",
        "### Deep-dive cases show how EasyICU separates output completion from interpretation",
        "",
        "Figure 6 should include two representative cases. The first is t04 lactate-mortality association, which shows a relatively clean association-analysis workflow and demonstrates how a figure, model summary and manuscript statement are linked to source evidence. The second is t08 vasopressor selection-bias audit, which is intentionally more cautionary. It demonstrates that a completed output may still require interpretation of warnings, repair history and design assumptions. Together, the cases show why EasyICU reports evidence and audit trails rather than presenting an agent-generated plot as a final clinical conclusion.",
        "",
        "## Discussion",
        "",
        "EasyICU provides a bridge between ICU data harmonisation and agent-assisted scientific analysis. Its main value is the coupling of clinical concept metadata, previewable cohort construction and evidence-bound execution. This coupling matters because most errors in ICU database research are not purely syntactic. They arise from ambiguous time windows, hidden unit conversions, source-specific measurement practice, missing components and unsupported interpretation. EasyICU makes these choices visible before and after analysis.",
        "",
        "The v15 benchmark supports this design. Clean completion across 60 audited task-arm cells shows that the final workflow can produce complete outputs across diverse ICU tasks. The 230 repair/fallback events show that auditability is essential: a final table or figure should not be reported without information about the assistance required to produce it. In a tool paper, this is a strength rather than a weakness, because it makes the boundary between model generation, deterministic repair and human interpretation explicit.",
        "",
        "The research-agent component should not be described as a fully autonomous clinical scientist. It is better described as an auditable assistant embedded inside a clinically aware runtime. It can accelerate repetitive analysis steps, generate scripts and draft evidence-linked text, but clinical interpretation, study design and submission decisions remain human responsibilities.",
        "",
        "Several limitations remain. EasyICU depends on the quality of the local database preparation and source mapping. The current final benchmark was run on one primary local coding model and one audited task registry, so broader model comparisons and external site replications are needed. The benchmark tasks include association and prediction examples, but they do not establish causal treatment effects. Finally, publication-quality figure design still requires author review even when the underlying data and audit status are complete.",
        "",
        "## Online Methods",
        "",
        "### Database preparation and clinical concept mapping",
        "",
        "EasyICU prepares public ICU databases through database-specific adapters and a shared clinical concept layer. Source tables are mapped to standard concepts, including demographics, ICU stay structure, vital signs, laboratory measurements, interventions, organ-support variables, severity-score components and outcomes. Each concept records source paths, units, time fields and extraction rules when available. The v16 manuscript should retain the v15 details for database versions, preprocessing assumptions and cohort denominators.",
        "",
        "### Preview-before-commit cohort construction",
        "",
        "The user-facing workflow separates exploration from export. Before committing a cohort, EasyICU displays source configuration, selected concepts, inclusion and exclusion filters, feature previews, patient-level time series, missingness summaries and cross-database distributions. This design reduces silent failures because users can inspect whether an extracted variable behaves as expected before downstream analysis scripts are generated.",
        "",
        "### SOFA-2 implementation and cross-database translation",
        "",
        "SOFA-2 is implemented as a component-wise clinical rule translation. The workflow records which component rules can be activated in each database and how missing source components affect score interpretation. The manuscript should preserve the v15 SOFA-2 rule table and component-activation results, because they demonstrate the cross-database method independent of the agent benchmark.",
        "",
        "### ResearchContext construction",
        "",
        "For agentic analysis, EasyICU converts an exported cohort into a ResearchContext. The context includes variable names, roles, likely units, time windows, missingness profiles, candidate outcomes, candidate exposures, allowed aggregations and ICU-specific warnings. This context is passed to the planner and coder so that generated code is grounded in clinical metadata rather than only column names.",
        "",
        "### Deterministic execution, validation and evidence binding",
        "",
        "Generated scripts are executed by a runtime supervisor. Outputs are inspected for expected files, metrics, table schemas and figure availability. When known failure modes occur, deterministic repair or fallback code may be applied and is recorded as a repair/fallback event. Every accepted artifact is registered in the EvidenceStore with a path, producer, metadata and SHA-256 hash. Manuscript scaffolds are filtered so that result-like sentences require evidence references before binding.",
        "",
        "### Fifteen-task benchmark protocol",
        "",
        "The final v15 benchmark used 15 ICU analysis tasks and four context/preference arms. A task-arm cell was marked clean_ok only when execution completed, expected artifacts were present, key metrics were extractable and no unresolved contract-level failure remained. Metric-sanity review checked the extracted metrics for obvious range or schema problems. Repair/fallback events were counted separately from final acceptance.",
        "",
        "### Figure and table curation",
        "",
        "Final publication candidates were copied or rebuilt into a curated publication package. Source maps link each curated figure and table back to the original evidence location. Figure 5 and Table 2 use the final audit CSV files rather than manually copied numbers. Figure 6 should be redrawn from audited task evidence for visual clarity, while retaining the source-map linkage.",
        "",
        "## Data availability",
        "",
        "The source ICU databases are available from their original data-use-controlled repositories. EasyICU-derived cohorts and benchmark artifacts are stored locally under the final run and final audit directories. Public release should include only artifacts permitted by the source database licenses and data-use agreements.",
        "",
        "## Code availability",
        "",
        "The EasyICU codebase contains the platform and research-agent modules used for this manuscript. The current build recorded the repository commit in the v16 package metadata. The author should confirm the final public commit before submission.",
        "",
        "## Acknowledgements",
        "",
        "[To be completed by the authors.]",
        "",
        "## Author contributions",
        "",
        "[To be completed by the authors.]",
        "",
        "## Competing interests",
        "",
        "[To be completed by the authors.]",
        "",
        "## References",
        "",
        "Keep and manually check the reference list from the v15 Nature Methods Word draft. No new references were fabricated during this generation step.",
        "",
        "## Figure legends",
        "",
        "### Figure 1. EasyICU system architecture for cross-database ICU research.",
        "EasyICU separates ICU data foundations, a safe analytical runtime, agent orchestration and scientific discovery outputs. The figure should show how database adapters, concept mappings, preview panels, cohort exports, ResearchContext objects, deterministic execution and evidence-bound writing connect into one workflow.",
        "",
        "### Figure 2. Preview-before-commit cohort construction workflow.",
        "Users inspect database configuration, concept selection, cohort filters, feature previews, patient-level trajectories, missingness and cross-database distributions before committing an analysis-ready export.",
        "",
        "### Figure 3. SOFA-2 cross-database clinical-rule translation.",
        "Component-level availability and activation are shown across databases, highlighting where missing components can affect score interpretation.",
        "",
        "### Figure 4. Evidence-bound research-agent runtime.",
        "EasyICU converts exported cohorts into ResearchContext metadata, plans analyses, executes generated code, validates outputs, registers evidence and binds manuscript statements to registered artifacts.",
        "",
        "### Figure 5. Fifteen-task four-arm benchmark of the EasyICU research-agent layer.",
        f"The benchmark included {len(matrix)} task-arm cells across 15 ICU analysis tasks and four settings. All {clean} cells reached clean completion in the final audit. Panels should show the acceptance heatmap, repair burden, task-band summaries and context/preference decomposition.",
        "",
        "### Figure 6. Deep-dive evidence trails for lactate association and vasopressor selection-bias audit.",
        "Two v15 cases show how EasyICU links final outputs to evidence, warnings, repair history and cautious interpretation.",
        "",
        "## Table legends",
        "",
        "### Table 1. Source-cohort characteristics across six public ICU databases.",
        "This table should be retained from the v15 Nature Methods draft and reviewed against the final cohort export logs.",
        "",
        "### Table 2. Fifteen-task benchmark summary with per-arm status and repair burden.",
        "This table is generated from the final v15 audit. Clean completion is reported together with repair/fallback counts to avoid overstating unassisted model performance.",
        "",
        "## Table 2 preview",
        "",
    ])
    fields = ["Task", "Plain-language task", "Task band", "aware status", "aware repairs", "aware_no_pref status", "aware_no_pref repairs", "naive_with_pref status", "naive_with_pref repairs", "naive status", "naive repairs", "Total repairs"]
    lines.extend(csv_markdown_table(table_rows, fields))
    return lines


def build_supp_lines(ctx: dict[str, object]) -> list[str]:
    matrix = ctx["matrix"]
    repair = ctx["repair"]
    metrics = ctx["metrics"]
    task_summary = ctx["task_summary"]
    figure_source = ctx["figure_source"]
    table_source = ctx["table_source"]
    repairs_by_phase = Counter(r.get("phase") or "unknown" for r in repair)
    repairs_by_task = Counter(r.get("task_key") or "unknown" for r in repair)
    repairs_by_arm = Counter(r.get("arm") or "unknown" for r in repair)
    clean = sum(1 for r in matrix if r.get("status") == "clean_ok")
    manual_flags = sum(1 for r in metrics if str(r.get("manual_review_required") or "").lower() == "true")
    lines = [
        "# Supplementary Information: EasyICU Nature Methods v16 FULL draft",
        "",
        "This supplementary file is generated as a full v16 author-review draft. It preserves the intended v15 supplementary structure conceptually while updating the final benchmark methods and tables to the 15-task four-arm audit.",
        "",
        "## Supplementary Methods M1. EasyICU platform setup and source-cohort preparation",
        "",
        "EasyICU prepares source cohorts from public ICU databases through database-specific adapters and a shared clinical concept layer. The author should merge the detailed database-version, extraction and cohort-denominator text from the v15 supplementary Word document into this section. The v16 manuscript keeps this material because the platform layer is the basis of the Nature Methods tool claim.",
        "",
        "## Supplementary Methods M2. SOFA-2 rule translation and preview workflow",
        "",
        "SOFA-2 component rules are translated through EasyICU clinical concepts. Component availability, score construction and missingness-sensitive interpretation should be reported for each database. Preview-before-commit panels allow users to inspect source mappings, feature rows, time series and missingness before export.",
        "",
        "## Supplementary Methods M3. Fifteen-task four-arm research-agent benchmark",
        "",
        "The v15 benchmark tested the EasyICU research-agent layer on 15 ICU analysis tasks. The task registry included descriptive summaries, stratified outcome summaries, severity-score correlation, association modelling, KDIGO sensitivity analysis, phenotype clustering, prediction performance, selection-bias audit, SOFA-zero artefact audit, missing-data robustness, length-of-stay summaries, age-stratified mortality, vital-sign summaries, creatinine trajectory/KDIGO analysis and norepinephrine-equivalent dose-response analysis.",
        "",
        "Each task was evaluated under four settings: ICU-aware with preferences, ICU-aware without preferences, generic context with preferences and generic context without preferences. The ICU-aware conditions supplied EasyICU ResearchContext metadata such as variable roles, units, time windows, allowed aggregations, missingness profiles and ICU-specific pitfalls. The generic conditions used reduced context analogous to a general data-agent prompt. The preference factor tested whether explicit user preferences changed completion or repair burden.",
        "",
        "A task-arm cell was marked clean_ok when execution completed, expected artifacts were present, key metrics were extractable and no unresolved contract-level failure remained. Repair/fallback events were counted separately from final acceptance. This distinction is important: clean completion indicates the final audited package is usable, whereas repair burden indicates how much deterministic assistance was required.",
        "",
        "## Supplementary Results R1. Final benchmark audit summary",
        "",
        f"The final audit included {len(matrix)} task-arm cells. All {clean} cells reached clean completion. Metric-sanity checks covered {len(metrics)} cells and produced {manual_flags} manual range-review flags. The audit recorded {len(repair)} repair or fallback events. Curated publication source maps listed {len(figure_source)} figure records and {len(table_source)} table records.",
        "",
        "### Repair/fallback events by arm",
        "",
    ]
    for arm in ARM_ORDER:
        lines.append(f"- {ARM_LABEL[arm]}: {repairs_by_arm.get(arm, 0)}")
    lines.extend(["", "### Repair/fallback events by phase", ""])
    for phase, count in repairs_by_phase.most_common():
        lines.append(f"- {phase}: {count}")
    lines.extend(["", "### Repair/fallback events by task", ""])
    for task, count in sorted(repairs_by_task.items(), key=lambda kv: task_num(kv[0])):
        lines.append(f"- {task}: {count}")
    lines.extend([
        "",
        "## Supplementary Tables",
        "",
        "### Supplementary Table S1. Database source summary",
        "Retain from the v15 supplementary document and verify against final EasyICU cohort exports.",
        "",
        "### Supplementary Table S2. Clinical concept mapping summary",
        "Retain from the v15 supplementary document and update only if source mappings changed.",
        "",
        "### Supplementary Table S3. Preview-before-commit panel definitions",
        "Retain or update from the v15 supplementary document.",
        "",
        "### Supplementary Table S4. SOFA-2 component rule translation",
        "Retain from the v15 supplementary document and cross-check component labels.",
        "",
        "### Supplementary Table S5. SOFA-2 availability and activation summary",
        "Retain from the v15 supplementary document.",
        "",
        "### Supplementary Table S10. Per-run acceptance summary",
        "Use `tables/S10_per_run_acceptance_summary.csv`. It contains one row per task-arm cell.",
        "",
        "### Supplementary Table S11. Context-ablation audit",
        "Use `tables/S11_context_ablation_audit.csv`. It adds ICU-context and user-preference indicators to each task-arm cell.",
        "",
        "### Supplementary Table S12. Fifteen-task numerical-results matrix",
        "Use `tables/S12_fifteen_task_numerical_results_matrix.csv`. It summarises task-level numerical outputs from the final audit and should be treated as validation evidence, not as new causal clinical findings.",
        "",
        "## Supplementary Figure guidance",
        "",
        "Supplementary figures should include source-cohort diagnostics, SOFA-2 component availability details, representative preview panels and benchmark audit details that are too granular for the main text. Figures that were visually acceptable but too simple for main text can be kept in the supplement. Figures marked as requiring redraw should not be promoted to the main manuscript without manual revision.",
        "",
        "## Supplementary audit tables embedded preview",
        "",
        "### S10 first ten rows",
        "",
    ])
    s10_fields = ["task_key", "task_label", "arm", "arm_label", "difficulty_band", "acceptance_status", "repair_events", "manual_metric_flags"]
    s10_path = Path(ctx["base"]) / "tables" / "S10_per_run_acceptance_summary.csv"
    s10_rows = read_csv(s10_path)
    lines.extend(csv_markdown_table(s10_rows, s10_fields, max_rows=10))
    lines.extend(["", "### S12 task-level metric summary", ""])
    if task_summary:
        fields = list(task_summary[0].keys())[:8]
        lines.extend(csv_markdown_table(task_summary, fields, max_rows=15))
    return lines


def build_checklist(ctx: dict[str, object], out_dir: Path) -> list[str]:
    return [
        "# EasyICU v16 FULL manuscript author checklist",
        "",
        "## Must check before submission",
        "",
        "- **Author list**: replace placeholder author list in the main manuscript.",
        "- **References**: paste and manually check the reference list from the v15 Nature Methods Word draft.",
        "- **Table 1**: merge the exact source-cohort numbers from the v15 official draft or final cohort export logs.",
        "- **Figure 1-4**: verify the existing architecture, workflow, SOFA-2 and agent-framework figures match the final text.",
        "- **Figure 5**: draw from `source_data/figure5_source_data.csv`.",
        "- **Figure 6**: redraw from audited t04 and t08 evidence; keep causal language out of the caption.",
        "- **Supplement**: merge old S1-S9 material with the new S10-S12 update.",
        "- **Overclaim check**: avoid fully autonomous, causal treatment effect, or mortality effect language unless explicitly supported.",
        "- **Data-use check**: ensure no protected patient-level data are exported into submission files.",
        "",
        "## Generated file roots",
        "",
        f"- **v16 package**: `{ctx['base']}`",
        f"- **final audit**: `{ctx['audit']}`",
        f"- **FULL package**: `{out_dir}`",
    ]


def build_file_index(out_dir: Path, ctx: dict[str, object]) -> list[dict[str, object]]:
    rows: list[dict[str, object]] = []
    for path in sorted(out_dir.rglob("*")):
        if path.is_file():
            rows.append({
                "relative_path": str(path.relative_to(out_dir)),
                "bytes": path.stat().st_size,
                "role": "generated_full_submission_file",
            })
    for rel in [
        "tables/Table2_15x4_benchmark_v16.csv",
        "tables/S10_per_run_acceptance_summary.csv",
        "tables/S11_context_ablation_audit.csv",
        "tables/S12_fifteen_task_numerical_results_matrix.csv",
        "source_data/figure5_source_data.csv",
        "figures/figure5_design.md",
        "figures/figure6_design.md",
    ]:
        p = Path(ctx["base"]) / rel
        rows.append({"relative_path": str(p), "bytes": p.stat().st_size if p.exists() else 0, "role": "supporting_v16_source_file"})
    return rows


def build(args: argparse.Namespace) -> None:
    ctx = load_context(args)
    base = Path(ctx["base"])
    out_dir = base / "FULL_submission_package_20260511"
    out_dir.mkdir(parents=True, exist_ok=True)
    main_lines = build_main_lines(ctx)
    supp_lines = build_supp_lines(ctx)
    write_text(out_dir / "EasyICU_main_manuscript_NatureMethods_v16_FULL_20260511.md", main_lines)
    markdown_to_docx(main_lines, out_dir / "EasyICU_main_manuscript_NatureMethods_v16_FULL_20260511.docx")
    write_text(out_dir / "EasyICU_supplementary_materials_NatureMethods_v16_FULL_20260511.md", supp_lines)
    markdown_to_docx(supp_lines, out_dir / "EasyICU_supplementary_materials_NatureMethods_v16_FULL_20260511.docx")
    checklist = build_checklist(ctx, out_dir)
    write_text(out_dir / "AUTHOR_FINAL_CHECKLIST.md", checklist)
    readme = [
        "# EasyICU Nature Methods v16 FULL submission package",
        "",
        "This folder contains a full author-review draft for the EasyICU Nature Methods tool paper.",
        "",
        "## Main files",
        "",
        "- `EasyICU_main_manuscript_NatureMethods_v16_FULL_20260511.docx`",
        "- `EasyICU_main_manuscript_NatureMethods_v16_FULL_20260511.md`",
        "- `EasyICU_supplementary_materials_NatureMethods_v16_FULL_20260511.docx`",
        "- `EasyICU_supplementary_materials_NatureMethods_v16_FULL_20260511.md`",
        "- `AUTHOR_FINAL_CHECKLIST.md`",
        "- `SUBMISSION_FILE_INDEX.csv`",
        "",
        "## Important framing",
        "",
        "This is an EasyICU tool/methods paper. The v15 benchmark is a validation module, not the whole paper.",
    ]
    write_text(out_dir / "00_README_START_HERE.md", readme)
    rows = build_file_index(out_dir, ctx)
    write_csv(out_dir / "SUBMISSION_FILE_INDEX.csv", rows, ["relative_path", "bytes", "role"])
    summary = {
        "generated_at": datetime.now().isoformat(),
        "out_dir": str(out_dir),
        "main_md_bytes": (out_dir / "EasyICU_main_manuscript_NatureMethods_v16_FULL_20260511.md").stat().st_size,
        "main_docx_bytes": (out_dir / "EasyICU_main_manuscript_NatureMethods_v16_FULL_20260511.docx").stat().st_size,
        "supp_md_bytes": (out_dir / "EasyICU_supplementary_materials_NatureMethods_v16_FULL_20260511.md").stat().st_size,
        "supp_docx_bytes": (out_dir / "EasyICU_supplementary_materials_NatureMethods_v16_FULL_20260511.docx").stat().st_size,
        "matrix_rows": len(ctx["matrix"]),
        "clean_ok": sum(1 for r in ctx["matrix"] if r.get("status") == "clean_ok"),
        "repair_events": len(ctx["repair"]),
        "metric_rows": len(ctx["metrics"]),
        "figure_source_rows": len(ctx["figure_source"]),
        "table_source_rows": len(ctx["table_source"]),
        "file_index_rows": len(rows),
    }
    (out_dir / "build_summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    print(out_dir)
    print(out_dir / "00_README_START_HERE.md")
    print(out_dir / "EasyICU_main_manuscript_NatureMethods_v16_FULL_20260511.docx")
    print(out_dir / "EasyICU_supplementary_materials_NatureMethods_v16_FULL_20260511.docx")
    print(out_dir / "build_summary.json")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--v16-dir", default="/Users/haibo/Documents/GitHub/easyicu写作/00_当前文章/01_当前正式稿/_nature_methods_v16_restructure_20260511")
    parser.add_argument("--writing-root", default="/Users/haibo/Documents/GitHub/easyicu写作/00_当前文章/01_当前正式稿")
    parser.add_argument("--audit-dir", default="/Users/haibo/Documents/GitHub/EASYICU/research_output/v15_experiments_20260509_1854_full/final_audit_20260510_60clean")
    build(parser.parse_args())
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
